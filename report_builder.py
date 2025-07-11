from docx import Document

def generate_docx(report_text, work_order_number):
    doc = Document()
    doc.add_heading("TEARDOWN/REJECT REPORT", level=1)

    doc.add_paragraph(report_text)
    doc.add_paragraph("\nAEROREPAIR SIGNATURE: ___________________________       DATE: __________")
    doc.add_paragraph("CUSTOMER SIGNATURE: _____________________________       DATE: __________")
    doc.add_paragraph("\nFAA CRS# R6CR725J")

    filename = f"Reject_Report_{work_order_number}.docx"
    doc.save(filename)
    print(f"Report saved as {filename}")
